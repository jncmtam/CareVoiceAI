#!/usr/bin/env python3
"""Tạo nhiều đơn thuốc mẫu .docx và in kết quả OCR mock."""

from __future__ import annotations

import json
import sys
from dataclasses import dataclass
from datetime import timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
BACKEND = ROOT / "backend"
OCR_TEST_DIR = BACKEND / "test" / "ocr"

sys.path.insert(0, str(BACKEND))

from app.integrations.vnpt.parsers.prescription import parse_ocr_payload  # noqa: E402
from app.utils.datetime import now_utc  # noqa: E402


@dataclass(frozen=True)
class PrescriptionTemplate:
    slug: str
    title: str
    lines: list[str]


PRESCRIPTIONS: list[PrescriptionTemplate] = [
    PrescriptionTemplate(
        slug="don_thuoc_chu_minh_tam",
        title="PHÒNG KHÁM NỘI TIẾT — BỆNH VIỆN DEMO CAREVOICE",
        lines=[
            "Mã hồ sơ: VC-2026-000001",
            "Bệnh nhân: Chu Minh Tâm",
            "Ngày sinh: 15/07/1998",
            "Giới tính: Nam",
            "SĐT: 0327628468",
            "Địa chỉ: TP.HCM",
            "Chẩn đoán: Đái tháo đường type 2, Tăng huyết áp",
            "",
            "Bác sĩ: BS. Lê Minh",
            "Khoa: Nội tiết",
            "Ngày kê đơn: 05/07/2026",
            "",
            "ĐƠN THUỐC:",
            "1. Metformin 500mg — 1 viên, uống 2 lần/ngày (sáng, tối), uống sau ăn.",
            "2. Amlodipine 5mg — 1 viên mỗi sáng, uống vào cùng một giờ mỗi ngày.",
            "",
            "Tái khám Nội tiết sau 14 ngày.",
            "",
            "Dặn dò: Uống thuốc đủ liều, không tự ý ngưng thuốc. Theo dõi đường huyết buổi sáng.",
            "Nếu chóng mặt, mệt bất thường hoặc đau ngực — gọi điều dưỡng ngay.",
            "",
            "Người kê đơn: BS. Lê Minh",
        ],
    ),
    PrescriptionTemplate(
        slug="don_thuoc_nguyen_thi_lan",
        title="PHÒNG KHÁM TIM MẠCH — BỆNH VIỆN DEMO CAREVOICE",
        lines=[
            "Mã hồ sơ: VC-2026-000002",
            "Bệnh nhân: Nguyễn Thị Lan",
            "Ngày sinh: 20/03/1958",
            "Giới tính: Nữ",
            "SĐT: 0901234567",
            "Địa chỉ: Quận 3, TP.HCM",
            "Chẩn đoán: Suy tim mạn, Rối loạn lipid máu",
            "",
            "Bác sĩ: BS. Phạm Hùng",
            "Khoa: Tim mạch",
            "Ngày kê đơn: 05/07/2026",
            "",
            "ĐƠN THUỐC:",
            "1. Bisoprolol 2.5mg — 1 viên mỗi sáng, uống trước ăn.",
            "2. Atorvastatin 20mg — 1 viên buổi tối, uống sau ăn.",
            "3. Furosemide 40mg — 1 viên buổi sáng, theo dõi cân nặng hàng ngày.",
            "",
            "Tái khám Tim mạch sau 21 ngày.",
            "",
            "Dặn dò: Hạn chế muối, theo dõi phù chân và cân nặng mỗi sáng.",
            "Nếu khó thở tăng hoặc đau ngực — liên hệ điều dưỡng ngay.",
            "",
            "Người kê đơn: BS. Phạm Hùng",
        ],
    ),
    PrescriptionTemplate(
        slug="don_thuoc_tran_van_binh",
        title="PHÒNG KHÁM HÔ HẤP — BỆNH VIỆN DEMO CAREVOICE",
        lines=[
            "Mã hồ sơ: VC-2026-000003",
            "Bệnh nhân: Trần Văn Bình",
            "Ngày sinh: 11/11/1965",
            "Giới tính: Nam",
            "SĐT: 0912345678",
            "Địa chỉ: Quận 7, TP.HCM",
            "Chẩn đoán: COPD, Hen phế quản",
            "",
            "Bác sĩ: BS. Nguyễn Thảo",
            "Khoa: Hô hấp",
            "Ngày kê đơn: 05/07/2026",
            "",
            "ĐƠN THUỐC:",
            "1. Salbutamol 100mcg — 2 nhát xịt khi khó thở, tối đa 4 lần/ngày.",
            "2. Budesonide 200mcg — 2 nhát xịt sáng và tối, súc miệng sau khi dùng.",
            "3. Theophylline 100mg — 1 viên 2 lần/ngày (sáng, chiều), uống sau ăn.",
            "",
            "Tái khám Hô hấp sau 10 ngày.",
            "",
            "Dặn dò: Không hút thuốc, tránh khói bụi. Mang theo thuốc xịt khi đi ra ngoài.",
            "",
            "Người kê đơn: BS. Nguyễn Thảo",
        ],
    ),
    PrescriptionTemplate(
        slug="don_thuoc_pham_thi_hoa",
        title="PHÒNG KHÁM CƠ XƯƠNG KHỚP — BỆNH VIỆN DEMO CAREVOICE",
        lines=[
            "Mã hồ sơ: VC-2026-000004",
            "Bệnh nhân: Phạm Thị Hoa",
            "Ngày sinh: 02/05/1952",
            "Giới tính: Nữ",
            "SĐT: 0934567890",
            "Địa chỉ: Bình Thạnh, TP.HCM",
            "Chẩn đoán: Loãng xương, Viêm khớp dạng thấp",
            "",
            "Bác sĩ: BS. Trần Quốc",
            "Khoa: Cơ xương khớp",
            "Ngày kê đơn: 05/07/2026",
            "",
            "ĐƠN THUỐC:",
            "1. Calcium 500mg — 1 viên 2 lần/ngày (sáng, tối), uống sau ăn.",
            "2. Vitamin D3 1000IU — 1 viên mỗi sáng.",
            "3. Methotrexate 7.5mg — uống 1 lần/tuần vào thứ Bảy, theo dõi men gan.",
            "",
            "Tái khám Cơ xương khớp sau 30 ngày.",
            "",
            "Dặn dò: Tập vận động nhẹ, tránh ngã. Không tự ý ngưng Methotrexate.",
            "",
            "Người kê đơn: BS. Trần Quốc",
        ],
    ),
    PrescriptionTemplate(
        slug="don_thuoc_le_van_duc",
        title="PHÒNG KHÁM THẬN — BỆNH VIỆN DEMO CAREVOICE",
        lines=[
            "Mã hồ sơ: VC-2026-000005",
            "Bệnh nhân: Lê Văn Đức",
            "Ngày sinh: 08/09/1970",
            "Giới tính: Nam",
            "SĐT: 0978123456",
            "Địa chỉ: Thủ Đức, TP.HCM",
            "Chẩn đoán: Suy thận mạn giai đoạn 3, Gout",
            "",
            "Bác sĩ: BS. Hoàng An",
            "Khoa: Thận",
            "Ngày kê đơn: 05/07/2026",
            "",
            "ĐƠN THUỐC:",
            "1. Losartan 50mg — 1 viên mỗi sáng, uống sau ăn.",
            "2. Allopurinol 100mg — 1 viên buổi tối, uống nhiều nước.",
            "3. Sodium bicarbonate 500mg — 1 viên 2 lần/ngày (sáng, tối).",
            "",
            "Tái khám Thận sau 14 ngày.",
            "",
            "Dặn dò: Hạn chế thịt đỏ và bia rượu. Theo dõi huyết áp và cân nặng hàng ngày.",
            "",
            "Người kê đơn: BS. Hoàng An",
        ],
    ),
    PrescriptionTemplate(
        slug="don_thuoc_vo_thi_mai",
        title="PHÒNG KHÁM NỘI THẦN KINH — BỆNH VIỆN DEMO CAREVOICE",
        lines=[
            "Mã hồ sơ: VC-2026-000006",
            "Bệnh nhân: Võ Thị Mai",
            "Ngày sinh: 14/12/1948",
            "Giới tính: Nữ",
            "SĐT: 0987654321",
            "Địa chỉ: Tân Bình, TP.HCM",
            "Chẩn đoán: Đột quỵ cũ, Tăng huyết áp",
            "",
            "Bác sĩ: BS. Lê Minh",
            "Khoa: Nội thần kinh",
            "Ngày kê đơn: 05/07/2026",
            "",
            "ĐƠN THUỐC:",
            "1. Clopidogrel 75mg — 1 viên mỗi sáng, uống sau ăn.",
            "2. Rosuvastatin 10mg — 1 viên buổi tối.",
            "3. Amlodipine 5mg — 1 viên mỗi sáng.",
            "4. Gabapentin 300mg — 1 viên 2 lần/ngày (sáng, tối), có thể gây buồn ngủ.",
            "",
            "Tái khám Nội thần kinh sau 28 ngày.",
            "",
            "Dặn dò: Không tự ý ngưng chống đông. Người nhà theo dõi nếu nói khó hoặc yếu liệt tay chân.",
            "",
            "Người kê đơn: BS. Lê Minh",
        ],
    ),
]


def build_prescription_docx(path: Path, template: PrescriptionTemplate) -> str:
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(template.title)
    run.bold = True
    run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("ĐƠN THUỐC").bold = True
    doc.add_paragraph("")

    raw_text = "\n".join(template.lines)
    for line in template.lines:
        doc.add_paragraph(line)

    doc.save(path)
    return raw_text


def run_ocr_preview(raw_text: str) -> dict:
    result = parse_ocr_payload(raw_text=raw_text)
    follow_up = result.draft_follow_up or {}
    if follow_up.get("appointment_at") is None and "ngày" in raw_text.lower():
        follow_up = {
            **follow_up,
            "appointment_at": (now_utc() + timedelta(days=14)).isoformat(),
        }
    return {
        "raw_text": result.raw_text,
        "draft_patient": result.draft_patient,
        "draft_medications": result.draft_medications,
        "draft_follow_up": follow_up,
        "instructions": result.instructions,
        "warnings": result.warnings,
    }


def main() -> None:
    OCR_TEST_DIR.mkdir(parents=True, exist_ok=True)
    index: list[dict] = []

    for template in PRESCRIPTIONS:
        docx_path = OCR_TEST_DIR / f"{template.slug}.docx"
        ocr_path = OCR_TEST_DIR / f"{template.slug}_ocr_result.json"
        raw_text = build_prescription_docx(docx_path, template)
        ocr = run_ocr_preview(raw_text)
        ocr_path.write_text(json.dumps(ocr, ensure_ascii=False, indent=2), encoding="utf-8")
        patient = (ocr.get("draft_patient") or {}).get("full_name")
        meds = len(ocr.get("draft_medications") or [])
        index.append(
            {
                "slug": template.slug,
                "docx": docx_path.name,
                "patient_name": patient,
                "medication_count": meds,
            }
        )
        print(f"✅ {docx_path.name} — {patient or '?'} — {meds} thuốc")

    index_path = OCR_TEST_DIR / "prescriptions_index.json"
    index_path.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
    print()
    print(f"Đã tạo {len(PRESCRIPTIONS)} đơn thuốc trong {OCR_TEST_DIR}")
    print(f"Index: {index_path}")


if __name__ == "__main__":
    main()